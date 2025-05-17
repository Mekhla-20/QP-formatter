import streamlit as st
import os
import zipfile
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Question Paper Formatter", layout="centered")
st.title("📝 Question Paper Formatter")

st.markdown("""
Upload multiple `.docx` files below and choose your formatting options. This tool supports English, Hindi, and Sanskrit documents.
""")

# Sidebar for formatting options
st.sidebar.header("⚙️ Formatting Options")
font_name = st.sidebar.selectbox("Font Name", ["Times New Roman", "Arial", "Calibri", "Mangal", "Kruti Dev"])
font_size = st.sidebar.slider("Font Size (pt)", 10, 18, 12)
line_spacing = st.sidebar.selectbox("Line Spacing", [1.0, 1.15, 1.5, 2.0])
margin_inch = st.sidebar.slider("Margins (inches)", 0.5, 2.0, 1.0)
bold_sections = st.sidebar.checkbox("Bold Section Headers", value=True)
auto_indent = st.sidebar.checkbox("Auto Indent Paragraphs", value=True)

add_header = st.sidebar.checkbox("Add Header/Footer", value=True)
school_name = st.sidebar.text_input("School Name", "Your School Name")
exam_name = st.sidebar.text_input("Exam Name", "Term 1 Examination")

uploaded_files = st.file_uploader("📂 Upload DOCX Files", type="docx", accept_multiple_files=True)

# Function to detect section headers
def is_section_header(text):
    section_keywords = ["section", "instructions", "general", "note"]
    return any(text.lower().strip().startswith(k) for k in section_keywords)

# Function to align marks at the end of a question line
def align_marks_right(paragraph):
    if "(" in paragraph.text and ")" in paragraph.text:
        text = paragraph.text
        parts = text.rsplit("(", 1)
        if len(parts) == 2 and ")" in parts[1]:
            question, marks = parts[0].strip(), "(" + parts[1].strip()
            paragraph.clear()
            run_q = paragraph.add_run(question + " ")
            run_m = paragraph.add_run(marks)
            run_q.font.name = run_m.font.name = font_name
            run_q._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run_m._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run_q.font.size = run_m.font.size = Pt(font_size)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Add page number field to footer
def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

# Function to format a single docx file
def format_docx(file, settings):
    doc = Document(file)

    # Page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(settings['margin'])
        section.bottom_margin = Inches(settings['margin'])
        section.left_margin = Inches(settings['margin'])
        section.right_margin = Inches(settings['margin'])

    # Header/Footer
    if settings['add_header']:
        for section in sections:
            header = section.header
            header.paragraphs[0].text = settings['school'] + " | " + settings['exam']
            footer = section.footer
            if footer.paragraphs:
                p = footer.paragraphs[0]
            else:
                p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_page_number(p)

    # Apply formatting to all paragraphs
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = settings['font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), settings['font'])
            run.font.size = Pt(settings['size'])

        para.paragraph_format.line_spacing = settings['spacing']

        if settings['indent']:
            para.paragraph_format.left_indent = Inches(0.25)

        # Detect and bold section headers
        if settings['bold_section'] and is_section_header(para.text):
            for run in para.runs:
                run.bold = True
            if "instruction" in para.text.lower():
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.left_indent = Inches(0.0)
                para.paragraph_format.first_line_indent = Inches(0.0)

        # Align marks right
        align_marks_right(para)

    return doc

if uploaded_files:
    formatted_files = []
    for uploaded_file in uploaded_files:
        formatted_doc = format_docx(uploaded_file, {
            'font': font_name,
            'size': font_size,
            'spacing': line_spacing,
            'margin': margin_inch,
            'add_header': add_header,
            'school': school_name,
            'exam': exam_name,
            'bold_section': bold_sections,
            'indent': auto_indent
        })

        out_stream = BytesIO()
        formatted_doc.save(out_stream)
        out_stream.seek(0)
        formatted_files.append((uploaded_file.name, out_stream))

    if formatted_files:
        # Create ZIP
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            for name, file_data in formatted_files:
                zip_file.writestr(name, file_data.read())
        zip_buffer.seek(0)

        st.success("✅ All files formatted successfully!")
        st.download_button("📦 Download All as ZIP", zip_buffer, file_name="Formatted_Questions.zip")
